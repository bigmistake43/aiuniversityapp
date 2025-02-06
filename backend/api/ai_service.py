from openai import OpenAI
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from io import BytesIO
import requests
from PIL import Image
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from django.conf import settings
from django.urls import reverse

client = OpenAI(
    api_key="Api-key", # Replace with your OpenAI API key
)

def generate_image_with_dalle(description):
    # Make an API request to DALL·E (replace with actual API call)
    dalle_api_url = "https://api.openai.com/v1/images/generations"
    headers = {
        "Authorization": "Bearer Api-key", # Replace with your OpenAI API key
        "Content-Type": "application/json"
    }
    payload = {
        "prompt": description,  # The description from GPT-4 response
        "n": 1,
        "size": "1024x1024"  # Adjust the image size as needed
    }
    response = requests.post(dalle_api_url, headers=headers, json=payload)
    image_url = response.json()['data'][0]['url']  # Assuming DALL·E returns a URL
    return image_url

def generate_chatbot_response(user_message, chat_type, pre_generated_prompts):

    base_system_prompt = """

                            You are a helpful assistant specializing in various topics including university recommendations, scholarships, courses, and location details.
                            Don't ask again to and response about the message based on the image you just given to user using all messages between you and user.
                            Also, You can generate images, pdf files, doc files.
                         """
    
    # Prepare messages for OpenAI API call
    messages = [{"role": "system", "content": base_system_prompt}]
    
    # Add pre-generated prompts to the system message
    for prompt in pre_generated_prompts:
        if isinstance(prompt, dict) and 'content' in prompt:
            prompt['content'] = str(prompt['content'])  # Ensure the content is a string
            messages.append({"role": "system", "content": prompt['content']})
        else:
            print("Invalid prompt:", prompt)

    messages.append({"role": "user", "content": user_message})
    

    try:
        # Make the OpenAI API request
        chat_completion = client.chat.completions.create(
            messages=messages,
            model="gpt-4o-mini" 
        )

        # Extract the response from OpenAI
        response_text = chat_completion.choices[0].message.content
        image_url = generate_image_with_dalle(response_text)

        response = requests.get(image_url)
        if response.status_code == 200:
            img = Image.open(BytesIO(response.content))
            rgb_img = img.convert('RGB')  # Convert to RGB (needed for JPEG format)

            # Save the image as JPG
            save_path = "media/uploads/generated_image.jpg"
            rgb_img.save(save_path, "JPEG")

            print(f"Image saved successfully as {save_path}")
        else:
            print("Failed to download the image. Check the URL.")

        html_response = markdown.markdown(response_text)

        # Initialize the result dictionary
        result = {'text_response': html_response, 'files': []}
        
        # Generate and save documents (PDF and DOCX)
        # pdf_file_path = generate_pdf(response_text)
        # docx_file_path = generate_docx(response_text)
        img_file_path = save_path
        
        # Convert file paths into URLs
        # pdf_file_url = settings.MEDIA_URL + pdf_file_path  # Use MEDIA_URL to create an accessible URL
        # docx_file_url = settings.MEDIA_URL + docx_file_path  # Use MEDIA_URL to create an accessible URL
        
        # Add file URLs to the result
        result['files'].append({
            'file_url': img_file_path,
            'file_name': 'generated_image.jpg'
        })
        # result['files'].append({
        #     'file_url': pdf_file_url,
        #     'file_name': 'generated_document.pdf'
        # })
        # result['files'].append({
        #     'file_url': docx_file_url,
        #     'file_name': 'generated_document.docx'
        # })
        
        # Check if OpenAI returned any files (assuming OpenAI might send file info in some format)
        # if 'file' in chat_completion:
        #     # Iterate over the files in the response (if any)
        #     for file_info in chat_completion['file']:
        #         file_url = file_info['url']  # The file URL returned by OpenAI (hypothetical structure)

        #         # Download the file
        #         file_data = requests.get(file_url).content
        #         file_name = file_info.get('filename', 'attachment')  # Use file name or generate one

        #         # Save the file in Django's default storage (you can use custom storage if needed)
        #         file = ContentFile(file_data)
        #         file_path = default_storage.save(f"uploads/{file_name}", file)

        #         # Convert the file path into an accessible URL
        #         file_url = settings.MEDIA_URL + file_path
                
        #         # Store the file's URL and name for returning to the user
        #         result['files'].append({
        #             'file_url': file_url,
        #             'file_name': file_name,
        #         })

        return result

    except Exception as e:
        print(f"Error occurred: {e}")
        return {"text_response": "Sorry, I couldn't generate a response at the moment. Please try again later.", "files": []}

def generate_pdf(text):
    """
    Generates a PDF from the given text and saves it to the Django media storage.
    """
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    # Set the font and start writing the content
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Generated Document:")
    c.drawString(100, 730, text)
    
    c.save()

    buffer.seek(0)
    
    # Save to Django storage
    pdf_path = default_storage.save("uploads/generated_document.pdf", ContentFile(buffer.read()))
    
    return pdf_path

def generate_docx(text):
    """
    Generates a DOCX document from the given text and saves it to Django media storage.
    """
    doc = Document()
    doc.add_heading('Generated Document', 0)
    
    # Add the response text as the document's content
    doc.add_paragraph(text)
    
    # Save the document to Django storage
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    docx_path = default_storage.save("uploads/generated_document.docx", ContentFile(docx_buffer.read()))
    
    return docx_path
